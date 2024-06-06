import os
import hashlib
import fitz
import requests
from urllib.parse import urljoin
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
import re
from pdfminer.high_level import extract_text
from docx import Document
from openpyxl import Workbook, load_workbook
from datetime import datetime
from time import sleep
import tkinter as tk
from tkinter import ttk
import threading

# Функция для генерации безопасного имени файла на основе хэш-суммы URL
def generate_safe_filename(url):
    return hashlib.md5(url.encode()).hexdigest()

# Функция для скачивания файлов
def download_file(url, local_filename, timeout=60, retries=5):
    attempt = 0
    while attempt < retries:
        try:
            with requests.get(url, stream=True, timeout=timeout) as r:
                r.raise_for_status()  # Проверка на успешный статус ответа
                with open(local_filename, 'wb') as f:
                    for chunk in r.iter_content(chunk_size=8192):
                        f.write(chunk)
            return local_filename
        except requests.exceptions.RequestException as e:
            log_error(f"Ошибка при скачивании {url}: {e}")
            attempt += 1
            sleep(5)  # Подождать 5 секунд перед следующей попыткой
    return None

# Функция для парсинга PDF
def parse_pdf(filepath):
    try:
        def is_scanned_page(page):
            """Проверяет, содержит ли страница текст."""
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if "lines" in block and block["lines"]:
                    return False
            return True

        def extract_text_patterns(page):
            """Extracts text patterns like multi-line headings."""
            patterns = []
            blocks = page.get_text("dict")["blocks"]

            for block in blocks:
                if "lines" not in block:
                    continue

                multi_line_text = []
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if text:
                            multi_line_text.append(text)
                
                # Combine multi-line text into a single string if there are multiple lines
                if multi_line_text:
                    combined_text = "  ".join(multi_line_text)
                    patterns.append(combined_text)
            
            return patterns

        def extract_text_between_keywords(pdf_path, start_keyword, end_keyword, page_number=1):
            doc = fitz.open(pdf_path)

            # Проверяем первые две страницы на наличие текста
            first_page = doc.load_page(0)
            second_page = doc.load_page(1)
            if is_scanned_page(first_page) and is_scanned_page(second_page):
                return "Документ является сканированным"

            extracted_text = []
            start_found = False

            # Загрузка указанной страницы (по умолчанию первой страницы)
            page = doc.load_page(page_number)
            blocks = page.get_text("dict")["blocks"]

            for block in blocks:
                if "lines" not in block:
                    continue

                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"]

                        # Проверяем наличие ключевого слова для начала
                        if start_keyword.lower() in text.lower():
                            start_found = True

                        # Если найдено ключевое слово начала, начинаем сохранять текст
                        if start_found:
                            extracted_text.append(text)

                        if end_keyword.lower() in text.lower():
                            result = "\n".join(extracted_text)

                            # Убираем последние две строки, если они содержат пустую строку и версию
                            result_lines = result.split("\n")
                            if len(result_lines) > 2 and not result_lines[-2].strip() and "версия" in result_lines[-1].lower():
                                result_lines = result_lines[:-2]
                            return " ".join(result_lines)
            if start_found == False:
                highlighted_words = extract_text_patterns(page)
                if highlighted_words:
                    return "Приближённая тема: " + "".join(highlighted_words)
                else:
                    return "Текст на второй странице не соответствует шаблону и не содержит выделенных слов"
            
            return " ".join(extracted_text)

        pdf_path = filepath  # Укажите путь к вашему PDF-файлу
        start_keyword = "ФГБОУ ВО «ВГТУ», ВГТУ"
        end_keyword = "версия"

        title = extract_text_between_keywords(pdf_path, start_keyword, end_keyword)

        with open(filepath, 'rb') as file:

            # Пытаемся здесь вычленить заголовок по первому шаблону
  
            # Здесь происходит парсинг док намбера
            text = extract_text(filepath)
            doc_number = None
            pattern = re.compile(r"\d+\.\d+\.\d+-\d+")
            matches = pattern.findall(text)

            unique_matches = set()
            for match in matches:
                unique_matches.add(match)

            for unique_match in unique_matches:
                print(unique_match)
                doc_number = unique_match
                
            return title, doc_number
    
    except Exception as e:
        log_error(f"Ошибка при парсинге PDF {filepath}: {e}")
        return 'Unknown', 'Unknown'

# Функция для проверки типа файла DOCX
def is_valid_word_file(filepath):
    try:
        doc = Document(filepath)
        return True
    except Exception as e:
        return False

# Функция для парсинга DOC
def parse_doc(filepath):
    try:
        if not is_valid_word_file(filepath):
            raise ValueError(f"Файл {filepath} не является валидным Word документом")

        doc = Document(filepath)
        title = doc.core_properties.title or "Unknown"

        potential_titles = []
        for paragraph in doc.paragraphs:
            if paragraph.runs and (any(run.bold for run in paragraph.runs) or paragraph.text.isupper()):
                potential_titles.append(paragraph.text)

        if potential_titles:
            title = potential_titles[0]

        doc_number = "Unknown"  # Нужно доработать под структуру документа
        
        return title, doc_number
    except Exception as e:
        log_error(f"Ошибка при парсинге DOC {filepath}: {e}")
        return 'Unknown', 'Unknown'

# Функция для парсинга сайта
def parse_website(base_url, timeout=60, retries=5):
    try:
        response = requests.get(base_url, timeout=timeout)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        documents = []
        unique_urls = set()
        current_header = None

        for element in soup.find_all(['h4', 'a']):
            if element.name == 'h4':
                current_header = element.text.strip()
            elif element.name == 'a' and ('href' in element.attrs) and (element['href'].endswith('.pdf') and ( str(element.text.strip()) != '' ) or element['href'].endswith('.doc') or element['href'].endswith('.docx')):
                file_url = urljoin(base_url, element['href'])
                file_display_name = element.text.strip()
                if file_url in unique_urls:
                    continue  # Пропустить дублирующийся URL
                unique_urls.add(file_url)

                file_name = os.path.basename(file_url)
                # file_display_name = element.text.strip()
                for attempt in range(retries):
                    try:
                        response = requests.head(file_url, timeout=timeout)
                        response.raise_for_status()
                        file_size = int(response.headers.get('Content-Length', 0))
                        documents.append({
                            'url': file_url,
                            'file_name': file_name,
                            'file_display_name': file_display_name,
                            'file_size': file_size,
                            'header': current_header
                        })
                        break
                    except requests.exceptions.RequestException as e:
                        log_error(f"Ошибка при обработке {file_url}: {e}")
                        if attempt < retries - 1:
                            sleep(5)  # Подождать 5 секунд перед следующей попыткой
                        else:
                            log_error(f"Не удалось получить размер файла для {file_url} после {retries} попыток.")
        return documents
    except requests.exceptions.RequestException as e:
        log_error(f"Ошибка при парсинге сайта {base_url}: {e}")
        return []

# Функция для проверки обновлений
def check_updates(documents, saved_data):
    current_files = {doc['url']: doc for doc in documents}
    saved_files = {doc['url']: doc for doc in saved_data}

    new_docs = [doc for doc in documents if doc['url'] not in saved_files]
    removed_docs = [doc for doc in saved_data if doc['url'] not in current_files]

    return new_docs, removed_docs

# Функция для сохранения данных в Excel
def save_to_excel(data, filename):
    wb = Workbook()
    ws = wb.active
    ws.append(['URL', 'Название файла', 'Отображаемое имя файла', 'Размер (байт)', 'Название документа', 'Номер документа', 'Дата'])

    current_header = None
    for doc in data:
        if doc['header'] != current_header:
            current_header = doc['header']
            ws.append([current_header])
        ws.append([
            doc['url'], doc['file_name'], doc['file_display_name'], doc['file_size'], 
            doc.get('title', 'Unknown'), doc.get('doc_number', 'Unknown'), 
            doc.get('timestamp', 'Unknown')
        ])

    wb.save(filename)

# Функция для обновления данных в GUI
def update_data():
    global documents, saved_data
    tree.delete(*tree.get_children())
    for doc in documents:
        tree.insert('', 'end', values=(doc['file_name'], doc.get('title', 'Unknown'), doc.get('doc_number', 'Unknown')))

# Функция для экспорта данных в Excel
def export_to_excel():
    global documents, excel_path
    save_to_excel(documents, excel_path)

# Функция для экспорта логов
def export_logs():
    global log_file
    with open(log_file, 'r') as file:
        logs = file.read()
    with open('logs_export.txt', 'w') as file:
        file.write(logs)

# Функция для логирования ошибок
def log_error(message):
    global log_text
    log_text.insert(tk.END, f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {message}\n")
    log_text.see(tk.END)
    with open(log_file, 'a') as log:
        log.write(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {message}\n")

# Функция для старта парсинга
def start_parsing():
    global parsing_thread, base_url, timeout, retries
    url = url_entry.get()
    if url:
        base_url = url
    parsing_thread = threading.Thread(target=parse_and_update)
    parsing_thread.start()
    progress_bar.start()

# Функция для остановки парсинга
def stop_parsing():
    global parsing_thread
    if parsing_thread.is_alive():
        # Невозможно безопасно остановить поток в Python, нужно предусмотреть флаг завершения
        pass

# Функция для выполнения парсинга и обновления данных
def check_updates(documents, saved_data):
    """
    Функция для сравнения новых данных с сохраненными данными и выявления изменений.

    Parameters:
        documents (list): Список новых документов.
        saved_data (list): Список сохраненных данных.

    Returns:
        tuple: Кортеж содержащий два списка - список новых документов и список удаленных документов.
    """
    new_docs = []
    removed_docs = []

    # Сравнение документов
    for doc in documents:
        if doc not in saved_data:
            new_docs.append(doc)

    for doc in saved_data:
        if doc not in documents:
            removed_docs.append(doc)

    return new_docs, removed_docs

# Ваш метод parse_and_update()
def parse_and_update():
    global documents, saved_data, timeout, retries

    # Получение новых данных с веб-сайта
    documents = parse_website(base_url, timeout=timeout, retries=retries)
    
    # Проверка обновлений
    new_docs, removed_docs = check_updates(documents, saved_data)

    # Обработка новых документов
    for doc in new_docs:
        safe_filename =  os.path.splitext(doc['file_name'])[1]
        local_filename = download_file(doc['url'], safe_filename, timeout=timeout, retries=retries)
        if local_filename:
            if local_filename.endswith('.pdf'):
                title, doc_number = parse_pdf(local_filename)
            elif local_filename.endswith('.doc') or local_filename.endswith('.docx'):
                title, doc_number = parse_doc(local_filename)
            else:
                title, doc_number = 'Unknown', 'Unknown'
            doc.update({
                'title': title,
                'doc_number': doc_number,
                'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
        else:
            log_error(f"Ошибка при скачивании {doc['url']}")
    
    # Если есть изменения, обновить GUI и остановить progress bar
    if new_docs or removed_docs:
        update_data()
        progress_bar.stop()
    
    # Обновление saved_data
    saved_data = documents.copy()


# Основная функция
def main():
    global documents, saved_data, base_url, excel_path, log_file, timeout, retries, tree, url_entry, progress_bar, parsing_thread, log_text

    base_url = 'https://cchgeu.ru/university/docs/'
    timeout = 60
    retries = 5
    log_file = 'parsing_errors.log'
    excel_path = 'documents.xlsx'
    documents = []
    saved_data = []

    if os.path.exists(excel_path):
        wb = load_workbook(excel_path)
        ws = wb.active
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] and row[1]:
                saved_data.append({
                    'url': row[0],
                    'file_name': row[1],
                    'file_display_name': row[2],
                    'file_size': row[3],
                    'title': row[4],
                    'doc_number': row[5],
                    'timestamp': row[6],
                    'header': row[7] if len(row) > 7 else None
                })


    root = tk.Tk()
    root.title("Document Parser")

    frame = ttk.Frame(root, padding="10")
    frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

    url_label = ttk.Label(frame, text="Введите URL:")
    url_label.grid(row=0, column=0, sticky=tk.W)
    url_entry = ttk.Entry(frame, width=50)
    url_entry.grid(row=0, column=1, sticky=(tk.W, tk.E))
    url_entry.insert(0, base_url)

    start_button = ttk.Button(frame, text="Начать парсинг", command=start_parsing)
    start_button.grid(row=1, column=0, sticky=tk.W)

    stop_button = ttk.Button(frame, text="Остановить парсинг", command=stop_parsing)
    stop_button.grid(row=1, column=1, sticky=tk.W)

    export_button = ttk.Button(frame, text="Сохарнить и экспортировать в Excel", command=export_to_excel)
    export_button.grid(row=1, column=2, sticky=tk.W)

    log_button = ttk.Button(frame, text="Экспортировать логи", command=export_logs)
    log_button.grid(row=1, column=3, sticky=tk.W)

    tree = ttk.Treeview(frame, columns=('file_name', 'title', 'doc_number'), show='headings')
    tree.heading('file_name', text='Имя файла')
    tree.heading('title', text='Название документа')
    tree.heading('doc_number', text='Номер документа')
    tree.grid(row=2, column=0, columnspan=4, sticky=(tk.W, tk.E, tk.N, tk.S))

    progress_bar = ttk.Progressbar(frame, mode='indeterminate')
    progress_bar.grid(row=3, column=0, columnspan=4, sticky=(tk.W, tk.E))

    log_text = tk.Text(frame, height=10, width=100)
    log_text.grid(row=4, column=0, columnspan=4, sticky=(tk.W, tk.E))

    root.columnconfigure(0, weight=1)
    root.rowconfigure(0, weight=1)

    frame.columnconfigure(1, weight=1)
    frame.rowconfigure(2, weight=1)

    root.mainloop()

if __name__ == "__main__":
    main()
